__author__ = 'posco'

from docx import Document
import xlrd
import os
import subprocess
import argparse

class HouseKeeping:
    def __init__(self, feedback_doc_name, marking_sheet_name, marker_name):
        print ("Feedback template: " + feedback_doc_name)
        print ("Marking sheet: " + marking_sheet_name)
        print ("Folder to check in: " + marker_name)

        self.feedback_doc_name = feedback_doc_name
        self.marker_name = marker_name
        #load student feedback form as a template
        self.feeback_document = Document(feedback_doc_name)
        #load my marking sheet 'PT' from workbook
        self.marking_sheet = xlrd.open_workbook(marking_sheet_name).sheet_by_name(marker_name)

    #do things
    def go(self):
        #username to firstname lastname map/dictionary
        self.name_map = {}
        self.construct_name_map()
        self.create_new_feedback_document()

		#probably won't work for Windows
    def unzip_submission(self, student_dir):
        #form unzip command
        cmd = 'unzip -d ' + student_dir + '/ ' + student_dir + '/*.zip'
        print cmd, '\n\n'
        sys_process = subprocess.Popen(cmd, shell=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        sys_process.wait()
        std_out = sys_process.stdout.read().strip()
        print std_out

    def create_new_feedback_document(self):
        marker_directory = os.path.dirname(os.path.realpath(__file__))+'/'+self.marker_name
        for student_dir, _, file in os.walk(marker_directory):
            student_dir_name = os.path.relpath(student_dir, marker_directory)

            #print student_dir
            if (student_dir_name is not '.') and (student_dir_name in self.name_map):
                student_name = self.name_map[student_dir_name][0] + ' ' + \
                               self.name_map[student_dir_name][1]
                self.write_student_name_to_document(student_dir, student_dir_name, student_name)

                #just do something extra
                self.unzip_submission(student_dir)


    def write_student_name_to_document(self, student_dir, student_dir_name, student_name):
        #default cell for student's firstname lastname
        filename = self.feedback_doc_name.replace('username', student_dir_name)
        self.feeback_document.tables[0].cell(1,0).text = student_name
        self.feeback_document.tables[0].cell(1,1).text = student_dir_name
        self.feeback_document.tables[0].cell(1,2).text = self.marker_name
        self.feeback_document.save(student_dir+'/'+filename)
        #print student_dir+'/'+filename

    def construct_name_map(self):
        username_index = 0
        is_constructing_name_map = False

        for i in range(self.marking_sheet.nrows):
            if is_constructing_name_map:
                username =  self.marking_sheet.row_values(i)[username_index]
                firstname = self.marking_sheet.row_values(i)[username_index-1]
                lastname =  self.marking_sheet.row_values(i)[username_index-2]
                self.name_map[username]=[firstname, lastname]

            elif self.marking_sheet.row_values(i).count('Username') is 1:
                username_index = self.marking_sheet.row_values(i).index('Username')
                is_constructing_name_map = True


parser = argparse.ArgumentParser(description='Housekeeping for 4001COMP marking.')
parser.add_argument('task', metavar='TASK', type=str, help='Task number (e.g. 1)')
parser.add_argument('initials', metavar='INITIALS', type=str, help='Marker initials (e.g. PH)')
args = parser.parse_args()

hk = HouseKeeping('feedback_username_task' + args.task + '.docx','4001COMP Marking 2014-15 CW1-T' + args.task + '.xlsx',args.initials)
hk.go()


